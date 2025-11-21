from __future__ import annotations
from pathlib import Path
from typing import List

import pikepdf
from PyPDF2 import PdfReader, PdfWriter
from pdf2image import convert_from_path
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4


def _update(job, status: str, message: str = "", progress: int | None = None):
    try:
        if job is None:
            return
        job.status = status
        job.message = message
        if progress is not None:
            job.progress = progress
    except Exception:
        pass


# ---------- Thumbnails ----------

def generate_pdf_thumbnails(src: Path, out_dir: Path, dpi: int = 110) -> List[Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    images = convert_from_path(str(src), dpi=dpi)
    paths: List[Path] = []
    for idx, img in enumerate(images, start=1):
        thumb = img.copy()
        thumb.thumbnail((320, 320))
        out_path = out_dir / f"{src.stem}_p{idx}.jpg"
        thumb.save(out_path, format="JPEG", quality=85)
        paths.append(out_path)
    return paths


# ---------- Merge ----------

def merge_pdfs_task(pdfs: List[str], out_path: Path, job_status):
    try:
        _update(job_status, "processing", "Merging PDFs", 5)
        with pikepdf.Pdf.new() as pdf_out:
            total = len(pdfs)
            for i, p in enumerate(pdfs, start=1):
                with pikepdf.Pdf.open(p) as src:
                    pdf_out.pages.extend(src.pages)
                _update(job_status, "processing", f"Merged {i}/{total}", int(5 + (i/total)*90))
            pdf_out.save(str(out_path))
        _update(job_status, "done", f"Saved to {out_path}", 100)
    except Exception as e:
        _update(job_status, "error", str(e))


# ---------- Split ----------

def parse_ranges(ranges: str, total_pages: int) -> List[tuple[int, int]]:
    if not ranges:
        return [(1, total_pages)]
    parts = [r.strip() for r in ranges.split(',') if r.strip()]
    result: List[tuple[int, int]] = []
    for part in parts:
        if '-' in part:
            a, b = part.split('-', 1)
            start = max(1, int(a))
            end = min(total_pages, int(b))
            if start <= end:
                result.append((start, end))
        else:
            i = int(part)
            if 1 <= i <= total_pages:
                result.append((i, i))
    return result


def split_pdf_task(src_pdf: str, ranges: str, out_dir: Path, job_status):
    try:
        _update(job_status, "processing", "Splitting", 5)
        reader = PdfReader(src_pdf)
        total = len(reader.pages)
        chunks = parse_ranges(ranges, total)
        for idx, (a, b) in enumerate(chunks, start=1):
            writer = PdfWriter()
            for page in range(a-1, b):
                writer.add_page(reader.pages[page])
            out_path = out_dir / f"split_{idx}_{a}-{b}.pdf"
            with open(out_path, 'wb') as f:
                writer.write(f)
            _update(job_status, "processing", f"Wrote part {idx}/{len(chunks)}", int(5 + (idx/len(chunks))*90))
        _update(job_status, "done", "Split complete", 100)
    except Exception as e:
        _update(job_status, "error", str(e))


# ---------- Reorder ----------

def reorder_pages_task(src_pdf: str, order: str, out_path: Path, job_status):
    try:
        reader = PdfReader(src_pdf)
        writer = PdfWriter()
        new_order = [int(x.strip()) for x in order.split(',') if x.strip()]
        for i in new_order:
            if 1 <= i <= len(reader.pages):
                writer.add_page(reader.pages[i-1])
        with open(out_path, 'wb') as f:
            writer.write(f)
        _update(job_status, "done", "Reordered", 100)
    except Exception as e:
        _update(job_status, "error", str(e))


# ---------- Rotate ----------

def rotate_pages_task(src_pdf: str, degrees: int, pages: str, out_path: Path, job_status):
    try:
        reader = PdfReader(src_pdf)
        writer = PdfWriter()
        to_rotate = set()
        if pages:
            to_rotate = {int(x.strip()) for x in pages.split(',') if x.strip()}
        for idx, page in enumerate(reader.pages, start=1):
            if not pages or idx in to_rotate:
                page.rotate(degrees)
            writer.add_page(page)
        with open(out_path, 'wb') as f:
            writer.write(f)
        _update(job_status, "done", "Rotated", 100)
    except Exception as e:
        _update(job_status, "error", str(e))


# ---------- Compress ----------

def compress_pdf_task(src_pdf: str, preset: str, out_path: Path, job_status):
    try:
        _update(job_status, "processing", "Compressing", 5)
        # Using pikepdf optimization; for stronger compression integrate Ghostscript externally.
        with pikepdf.open(src_pdf) as pdf:
            pdf.save(str(out_path), optimize_version=True, linearize=True)
        _update(job_status, "done", f"Compressed ({preset})", 100)
    except Exception as e:
        _update(job_status, "error", str(e))


# ---------- Convert ----------
from textwrap import wrap as textwrap

def convert_task(src_path: str, target: str, out_path: Path, job_status):
    try:
        _update(job_status, "processing", "Converting", 5)
        src = Path(src_path)
        target = target.lower()
        out_path.parent.mkdir(parents=True, exist_ok=True)

        if src.suffix.lower() == '.pdf' and target == 'docx':
            import pdfplumber
            from docx import Document
            doc = Document()
            with pdfplumber.open(str(src)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ''
                    doc.add_paragraph(text)
                    doc.add_page_break()
            doc.save(str(out_path))
        elif src.suffix.lower() == '.docx' and target == 'pdf':
            from docx import Document
            document = Document(str(src))
            c = canvas.Canvas(str(out_path), pagesize=letter)
            width, height = letter
            y = height - 60
            for para in document.paragraphs:
                c.setFont("Helvetica", 12)
                text = para.text
                for line in textwrap(text, 90):
                    c.drawString(40, y, line)
                    y -= 14
                    if y < 60:
                        c.showPage()
                        y = height - 60
            c.save()
        elif src.suffix.lower() in ('.jpg', '.jpeg', '.png') and target == 'pdf':
            img = Image.open(src)
            img = img.convert('RGB')
            img.save(str(out_path), "PDF", resolution=100.0)
        elif src.suffix.lower() == '.txt' and target == 'pdf':
            c = canvas.Canvas(str(out_path), pagesize=A4)
            width, height = A4
            y = height - 40
            with open(src, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    c.drawString(40, y, line.strip())
                    y -= 14
                    if y < 40:
                        c.showPage()
                        y = height - 40
            c.save()
        elif src.suffix.lower() == '.csv' and target == 'xlsx':
            import pandas as pd
            df = pd.read_csv(src)
            df.to_excel(str(out_path), index=False)
        elif src.suffix.lower() in ('.jpg', '.jpeg') and target == 'png':
            Image.open(src).save(str(out_path))
        elif src.suffix.lower() == '.png' and target == 'jpg':
            Image.open(src).convert('RGB').save(str(out_path))
        else:
            raise ValueError(f"Unsupported conversion: {src.suffix} -> {target}")

        _update(job_status, "done", "Converted", 100)
    except Exception as e:
        _update(job_status, "error", str(e))


# ---------- Edit (simple) ----------

def edit_to_docx_task(content: str, out_path: Path, job_status):
    try:
        from docx import Document
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        for para in content.split('\n'):
            doc.add_paragraph(para)
        doc.save(str(out_path))
        _update(job_status, "done", "Saved DOCX", 100)
    except Exception as e:
        _update(job_status, "error", str(e))


def edit_to_pdf_task(content: str, out_path: Path, job_status):
    try:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        c = canvas.Canvas(str(out_path), pagesize=letter)
        width, height = letter
        y = height - 60
        for line in content.split('\n'):
            c.setFont("Helvetica", 12)
            c.drawString(40, y, line)
            y -= 14
            if y < 60:
                c.showPage()
                y = height - 60
        c.save()
        _update(job_status, "done", "Saved PDF", 100)
    except Exception as e:
        _update(job_status, "error", str(e))
